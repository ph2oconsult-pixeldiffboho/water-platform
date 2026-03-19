"""
core/reporting/report_exporter.py

Two report formats:
  1. Executive Summary (1-2 pages) — key metrics, comparison table, recommendation
  2. Comprehensive Report — full sections, narrative, charts-as-tables, assumptions

Both available as PDF (reportlab) and Word (.docx).

Usage:
    from core.reporting.report_exporter import ReportExporter
    pdf_exec  = ReportExporter.to_pdf(report, mode="executive")
    pdf_full  = ReportExporter.to_pdf(report, mode="comprehensive")
    docx_exec = ReportExporter.to_docx(report, mode="executive")
    docx_full = ReportExporter.to_docx(report, mode="comprehensive")
"""
from __future__ import annotations
import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from core.reporting.report_engine import ReportObject

# ── Brand colours ──────────────────────────────────────────────────────────
BLUE_HEX   = "1F6AA5"
BLUE_LT    = "D5E8F0"
GREY_HEX   = "555555"
GREY_LT    = "F4F4F4"
RED_HEX    = "C0392B"
GREEN_HEX  = "27AE60"
ORANGE_HEX = "E67E22"


class ReportExporter:

    @staticmethod
    def to_pdf(report: ReportObject, mode: str = "comprehensive") -> bytes:
        if mode == "executive":
            return _pdf_executive(report)
        return _pdf_comprehensive(report)

    @staticmethod
    def to_docx(report: ReportObject, mode: str = "comprehensive") -> bytes:
        if mode == "executive":
            return _docx_executive(report)
        return _docx_comprehensive(report)


# ─────────────────────────────────────────────────────────────────────────────
# PDF HELPERS
# ─────────────────────────────────────────────────────────────────────────────

# ── Unicode → ReportLab XML markup ───────────────────────────────────────────
_SUB_MAP = {'₀':'0','₁':'1','₂':'2','₃':'3','₄':'4',
            '₅':'5','₆':'6','₇':'7','₈':'8','₉':'9'}
_SUP_MAP = {'⁰':'0','¹':'1','²':'2','³':'3','⁴':'4',
            '⁵':'5','⁶':'6','⁷':'7','⁸':'8','⁹':'9'}

def rl_safe(text: str) -> str:
    """
    Convert Unicode subscript/superscript digits to ReportLab XML markup
    so they render correctly with Helvetica in PDF output.
    e.g. "CO2e" -> "CO<sub>2</sub>e"
         "m2"   -> "m<super>2</super>"
    Also escapes & < > for ReportLab XML safety (except existing tags).
    """
    if not text:
        return text
    i, out = 0, []
    while i < len(text):
        c = text[i]
        if c in _SUB_MAP:
            run = []
            while i < len(text) and text[i] in _SUB_MAP:
                run.append(_SUB_MAP[text[i]]); i += 1
            out.append(f'<sub>{"".join(run)}</sub>')
        elif c in _SUP_MAP:
            run = []
            while i < len(text) and text[i] in _SUP_MAP:
                run.append(_SUP_MAP[text[i]]); i += 1
            out.append(f'<super>{"".join(run)}</super>')
        elif c == '²':   # U+00B2 superscript two
            out.append('<super>2</super>'); i += 1
        elif c == '³':   # U+00B3 superscript three
            out.append('<super>3</super>'); i += 1
        elif c == '°':
            out.append('°'); i += 1
        else:
            out.append(c); i += 1
    return ''.join(out)


def P(text, style):
    """Paragraph with rl_safe encoding applied."""
    from reportlab.platypus import Paragraph as _P
    return _P(rl_safe(str(text)), style)



def _pdf_styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib.units import mm

    rl_blue   = colors.HexColor(f"#{BLUE_HEX}")
    rl_lt     = colors.HexColor(f"#{BLUE_LT}")
    rl_grey   = colors.HexColor(f"#{GREY_HEX}")
    rl_ltgrey = colors.HexColor(f"#{GREY_LT}")

    base = getSampleStyleSheet()
    return {
        "title":    ParagraphStyle("rpt_title", parent=base["Normal"],
                       fontSize=22, fontName="Helvetica-Bold",
                       textColor=rl_blue, spaceAfter=4, spaceBefore=0),
        "subtitle": ParagraphStyle("rpt_subtitle", parent=base["Normal"],
                       fontSize=11, fontName="Helvetica", textColor=rl_grey,
                       spaceAfter=10),
        "h1":       ParagraphStyle("rpt_h1", parent=base["Normal"],
                       fontSize=13, fontName="Helvetica-Bold", textColor=rl_blue,
                       spaceBefore=14, spaceAfter=4),
        "h2":       ParagraphStyle("rpt_h2", parent=base["Normal"],
                       fontSize=10, fontName="Helvetica-Bold", textColor=rl_grey,
                       spaceBefore=8, spaceAfter=2),
        "body":     ParagraphStyle("rpt_body", parent=base["Normal"],
                       fontSize=9, fontName="Helvetica", leading=13,
                       spaceBefore=3, spaceAfter=3, alignment=TA_JUSTIFY),
        "bullet":   ParagraphStyle("rpt_bullet", parent=base["Normal"],
                       fontSize=9, fontName="Helvetica", leading=13,
                       spaceBefore=1, spaceAfter=1, leftIndent=12,
                       bulletIndent=4),
        "caption":  ParagraphStyle("rpt_caption", parent=base["Normal"],
                       fontSize=7.5, fontName="Helvetica-Oblique",
                       textColor=rl_grey, spaceBefore=2, spaceAfter=4),
        "disc":     ParagraphStyle("rpt_disc", parent=base["Normal"],
                       fontSize=7.5, fontName="Helvetica", textColor=rl_grey,
                       spaceBefore=4, spaceAfter=4, alignment=TA_JUSTIFY),
        "kpi_val":  ParagraphStyle("rpt_kpi_val", parent=base["Normal"],
                       fontSize=18, fontName="Helvetica-Bold",
                       textColor=rl_blue, spaceAfter=0, spaceBefore=0),
        "kpi_lbl":  ParagraphStyle("rpt_kpi_lbl", parent=base["Normal"],
                       fontSize=7, fontName="Helvetica", textColor=rl_grey,
                       spaceAfter=0, spaceBefore=0),
    }, {"blue": rl_blue, "lt": rl_lt, "grey": rl_grey, "ltgrey": rl_ltgrey}


def _pdf_tbl_style(colours, header=True, zebra=True):
    from reportlab.platypus import TableStyle
    from reportlab.lib import colors
    cmds = [
        ("FONTNAME",      (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE",      (0,0), (-1,-1), 8),
        ("TOPPADDING",    (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("LEFTPADDING",   (0,0), (-1,-1), 5),
        ("RIGHTPADDING",  (0,0), (-1,-1), 5),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("GRID",          (0,0), (-1,-1), 0.25, colors.HexColor("#CCCCCC")),
        ("ALIGN",         (1,0), (-1,-1), "RIGHT"),
        ("ALIGN",         (0,0), (0,-1), "LEFT"),
    ]
    if header:
        cmds += [
            ("BACKGROUND",  (0,0), (-1,0), colours["blue"]),
            ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
            ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,0), 8),
            ("ALIGN",       (0,0), (-1,0), "CENTER"),
        ]
    if zebra:
        cmds.append(("ROWBACKGROUNDS", (0,1), (-1,-1),
                     [colors.white, colours["ltgrey"]]))
    return TableStyle(cmds)


def _pdf_header_footer(canvas, doc, report, is_exec=False):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    w, h = A4
    canvas.saveState()
    # Header
    canvas.setFillColorRGB(0x1F/255, 0x6A/255, 0xA5/255)
    canvas.rect(0, h-22*mm, w, 22*mm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(18*mm, h-13*mm, report.project_name or "Treatment Technology Study")
    canvas.setFont("Helvetica", 8)
    rtype = "Executive Summary" if is_exec else "Detailed Report"
    canvas.drawRightString(w-18*mm, h-13*mm,
        f"{rtype} | {datetime.now().strftime('%d %b %Y')}")
    # Footer
    canvas.setFillColorRGB(0x55/255, 0x55/255, 0x55/255)
    canvas.setFont("Helvetica", 6.5)
    canvas.drawString(18*mm, 10*mm,
        "CONCEPT LEVEL — ±40% — NOT FOR PROCUREMENT OR FUNDING APPROVAL  |  "
        f"ph2o consulting  |  {report.plant_name or ''}")
    canvas.drawRightString(w-18*mm, 10*mm, f"Page {doc.page}")
    canvas.restoreState()


def _make_pdf_doc(buf, is_exec=False):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.lib.units import mm
    return SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=28*mm, bottomMargin=20*mm,
        leftMargin=18*mm, rightMargin=18*mm,
    )


def _chart_capex(chart_data, width, height=140, label_fmt=None):
    """Bar chart: CAPEX by scenario ($M)."""
    from reportlab.graphics.shapes import Drawing, String, Rect
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.legends import Legend
    from reportlab.lib import colors

    names  = chart_data.get("x", [])
    values = chart_data.get("y", [])
    if not names or not values:
        return None

    PALETTE = [colors.HexColor("#1F6AA5"), colors.HexColor("#27AE60"),
               colors.HexColor("#E67E22"), colors.HexColor("#8E44AD"),
               colors.HexColor("#C0392B")]

    d = Drawing(width, height)
    chart_w, chart_h = width - 60, height - 50

    bc = VerticalBarChart()
    bc.x, bc.y = 50, 35
    bc.width, bc.height = chart_w, chart_h
    bc.data = [[v] for v in values]
    bc.categoryAxis.categoryNames = [""]
    bc.valueAxis.valueMin = round(min(values) * 0.85, 1) if min(values) > 0 else 0
    mx = max(values)
    bc.valueAxis.valueMax = round(mx * 1.15 + 0.1, 1)
    step = round((bc.valueAxis.valueMax - bc.valueAxis.valueMin) / 4, 1)
    bc.valueAxis.valueStep = max(step, 0.1)
    bc.valueAxis.labels.fontSize = 7
    bc.valueAxis.labels.fontName = "Helvetica"
    bc.categoryAxis.labels.fontSize = 7
    bc.categoryAxis.labels.fontName = "Helvetica"
    bc.categoryAxis.visibleTicks = False
    bc.groupSpacing = 10
    bc.barSpacing = 3
    for i in range(len(values)):
        bc.bars[i].fillColor = PALETTE[i % len(PALETTE)]
        bc.bars[i].strokeColor = None
    d.add(bc)

    col_w = chart_w / len(values)
    for i, val in enumerate(values):
        bar_top = 35 + (val / (bc.valueAxis.valueMax or 1)) * chart_h
        lbl = label_fmt.format(val) if label_fmt else f"${val:.1f}M"
        d.add(String(50 + (i + 0.5) * col_w, bar_top + 3,
                     lbl, fontSize=7, fontName="Helvetica-Bold",
                     fillColor=PALETTE[i % len(PALETTE)], textAnchor="middle"))

    leg = Legend()
    leg.x, leg.y = 50, 10
    leg.dx, leg.dy = 8, 8
    leg.fontName, leg.fontSize = "Helvetica", 7
    leg.columnMaximum = len(names)
    leg.colorNamePairs = [(PALETTE[i % len(PALETTE)], n) for i, n in enumerate(names)]
    leg.alignment = "left"
    d.add(leg)
    return d


def _chart_stacked_h(labels, scenario_data, width, height=160, unit="k$/yr"):
    """Stacked horizontal bar chart."""
    from reportlab.graphics.shapes import Drawing, String, Rect
    from reportlab.lib import colors

    if not labels or not scenario_data:
        return None

    PALETTE = [colors.HexColor("#1F6AA5"), colors.HexColor("#27AE60"),
               colors.HexColor("#E67E22"), colors.HexColor("#8E44AD"),
               colors.HexColor("#C0392B"), colors.HexColor("#16A085"),
               colors.HexColor("#F39C12")]
    grey = colors.HexColor("#555555")

    scenarios = list(scenario_data.keys())
    n = len(scenarios)
    d = Drawing(width, height)

    LEFT = 95
    chart_w = width - LEFT - 15
    slot_h = (height - 35) / max(n, 1)
    bar_h = slot_h * 0.55

    totals = [sum(scenario_data[s].get(lb, 0) for lb in labels) for s in scenarios]
    max_t = max(totals) if totals else 1

    legend_seen = []
    for si, scen in enumerate(scenarios):
        y_bar = height - 25 - (si + 1) * slot_h + (slot_h - bar_h) / 2
        d.add(String(LEFT - 5, y_bar + bar_h * 0.3, scen,
                     fontSize=7, fontName="Helvetica-Bold",
                     fillColor=grey, textAnchor="end"))
        x = LEFT
        for li, lb in enumerate(labels):
            val = scenario_data[scen].get(lb, 0)
            if val <= 0:
                continue
            seg_w = (val / max_t) * chart_w
            c = PALETTE[li % len(PALETTE)]
            d.add(Rect(x, y_bar, seg_w, bar_h, fillColor=c, strokeColor=None))
            x += seg_w
            if lb not in legend_seen:
                legend_seen.append(lb)
        row_total = sum(scenario_data[scen].get(lb, 0) for lb in labels)
        d.add(String(x + 3, y_bar + bar_h * 0.3,
                     f"${row_total:.0f}k" if "k$" in unit else f"{row_total:.0f}",
                     fontSize=6.5, fontName="Helvetica", fillColor=grey))

    lx = LEFT
    for li, lb in enumerate(legend_seen[:7]):
        c = PALETTE[li % len(PALETTE)]
        d.add(Rect(lx, 4, 8, 8, fillColor=c, strokeColor=None))
        short = lb[:16]
        d.add(String(lx + 10, 5, short, fontSize=6, fontName="Helvetica", fillColor=grey))
        lx += len(short) * 4.2 + 16

    return d



def _pdf_hr(colours):
    from reportlab.platypus import HRFlowable
    return HRFlowable(width="100%", thickness=0.6,
                      color=colours["blue"], spaceAfter=4, spaceBefore=2)


def _render_dict_table(data: dict, styles, colours, W):
    """Render a {headers, rows} dict as a PDF Table."""
    from reportlab.platypus import Table, Paragraph
    from reportlab.lib.units import mm
    hdrs = data.get("headers", [])
    rows = data.get("rows", [])
    if not hdrs or not rows:
        return None
    n = len(hdrs)

    # ── OPEX breakdown table: many narrow columns — use smaller font + equal widths
    is_opex = any("%" in str(h) for h in hdrs)
    if is_opex and n > 6:
        col_w = [W * 0.18] + [W * 0.82 / (n - 1)] * (n - 1)
        cell_style = styles.get("caption") or styles["body"]
    elif n >= 2:
        col_w = [W * 0.28] + [W * 0.72 / (n-1)] * (n-1)
        cell_style = styles["body"]
    else:
        col_w = [W / n] * n
        cell_style = styles["body"]

    tbl_data = [[Paragraph(rl_safe(str(h)), styles["h2"]) for h in hdrs]]
    for r in rows:
        tbl_data.append([Paragraph(rl_safe(str(r.get(h, "—"))), cell_style) for h in hdrs])
    t = Table(tbl_data, colWidths=col_w, repeatRows=1)
    t.setStyle(_pdf_tbl_style(colours))
    return t


def _render_list_table(data: list, styles, colours, W, first_col_wide=True):
    """Render a list-of-dicts as a PDF Table (comparison style)."""
    from reportlab.platypus import Table, Paragraph
    if not data:
        return None
    hdrs = list(data[0].keys())
    n = len(hdrs)
    if first_col_wide and n >= 2:
        col_w = [W * 0.30] + [W * 0.70 / (n-1)] * (n-1)
    else:
        col_w = [W / n] * n
    tbl_data = [[Paragraph(rl_safe(str(h)), styles["h2"]) for h in hdrs]]
    for r in data:
        tbl_data.append([Paragraph(rl_safe(str(r.get(h, "—"))), styles["body"]) for h in hdrs])
    t = Table(tbl_data, colWidths=col_w, repeatRows=1)
    t.setStyle(_pdf_tbl_style(colours))
    return t


def _format_assumption_param(raw: str) -> str:
    """Convert raw assumption key to readable label."""
    label_map = {
        "electricity_per_kwh": "Electricity price ($/kWh)",
        "carbon_price_per_tonne": "Carbon price ($/tCO2e)",
        "discount_rate": "Discount rate",
        "analysis_period_years": "Analysis period (years)",
        "grid_emission_factor": "Grid emission factor (kg CO2e/kWh)",
        "aeration_tank_per_m3": "Aeration tank civil ($/m³)",
        "secondary_clarifier_per_m2": "Secondary clarifier ($/m2)",
        "blower_per_kw": "Blower system ($/kW)",
        "mbr_membrane_per_m2": "MBR membrane ($/m2)",
        "sludge_disposal_per_tds": "Sludge disposal ($/t DS)",
        "methanol_per_kg": "Methanol ($/kg)",
        "ferric_chloride_per_kg": "Ferric chloride ($/kg)",
    }
    # Handle nested keys like "capex_unit_costs → aeration_tank_per_m3"
    if "→" in raw:
        parts = raw.split("→")
        sub = parts[-1].strip()
        return label_map.get(sub, sub.replace("_", " ").title())
    return label_map.get(raw, raw.replace("_", " ").title())


# ─────────────────────────────────────────────────────────────────────────────
# EXECUTIVE SUMMARY PDF  (1–2 pages)
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_executive(report: ReportObject) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm, cm
    from reportlab.platypus import (
        Paragraph, Spacer, Table, TableStyle, KeepTogether,
    )
    from reportlab.lib import colors

    styles, colours = _pdf_styles()
    buf = io.BytesIO()
    doc = _make_pdf_doc(buf, is_exec=True)
    W = A4[0] - 36*mm

    def hf(canvas, doc):
        _pdf_header_footer(canvas, doc, report, is_exec=True)

    story = []

    # ── Title block ────────────────────────────────────────────────────────
    story.append(Paragraph(
        report.project_name or "Treatment Technology Comparison", styles["title"]))
    story.append(Paragraph(
        f"{report.plant_name or 'Wastewater Treatment'} — Executive Summary",
        styles["subtitle"]))
    story.append(Paragraph(
        f"Prepared by: {report.prepared_by or 'ph2o consulting'}  |  "
        f"Date: {datetime.now().strftime('%B %Y')}  |  "
        f"Scenarios evaluated: {len(report.scenario_names)}",
        styles["caption"]))
    story.append(_pdf_hr(colours))
    story.append(Spacer(1, 4))

    # ── Study description ──────────────────────────────────────────────────
    story.append(Paragraph("Study Overview", styles["h1"]))
    story.append(Paragraph(
        f"This executive summary presents the results of a concept-stage treatment "
        f"technology comparison study for <b>{report.plant_name or 'the facility'}</b>. "
        f"A total of <b>{len(report.scenario_names)} treatment scenarios</b> were evaluated "
        f"across capital cost, operating cost, lifecycle cost, energy consumption, carbon "
        f"footprint, and risk. All estimates are concept-level (±40%) and are intended for "
        f"comparative purposes only.",
        styles["body"]))
    story.append(Spacer(1, 4))

    # ── Scenarios overview table ───────────────────────────────────────────
    story.append(Paragraph("Scenarios Evaluated", styles["h1"]))
    story.append(_pdf_hr(colours))
    scen_rows = [["Scenario", "Technology", "Key Characteristic"]]
    tech_notes = {
        "bnr": "BNR with Secondary Clarifiers",
        "granular_sludge": "Aerobic Granular Sludge (Nereda®)",
        "mabr_bnr": "MABR retrofit into BNR aerobic zone",
        "bnr_mbr": "BNR basins + Membrane Separation",
        "ifas_mbbr": "IFAS carrier media retrofit",
        "anmbr": "Anaerobic MBR",
    }
    # Build a lookup from scenario_name -> tech_code via comparison_table data
    # The comparison_table has scenario names as column headers but not tech codes.
    # Use the report's decision object or scenario objects if available.
    _scenario_tech_map = {}
    _scenario_chars = {
        "bnr":             ("BNR (Activated Sludge)",        "Conventional BNR + secondary clarifiers"),
        "granular_sludge": ("Aerobic Granular Sludge",       "SBR-based granular sludge, no clarifiers"),
        "mabr_bnr":        ("MABR + BNR",                    "Membrane-aerated biofilm in BNR aerobic zone"),
        "bnr_mbr":         ("BNR + MBR Hybrid",              "BNR basins with membrane separation"),
        "ifas_mbbr":       ("IFAS / MBBR",                   "Integrated fixed-film activated sludge"),
        "anmbr":           ("Anaerobic MBR",                 "Anaerobic membrane bioreactor"),
        "mob":             ("MOB Biofilm",                    "Microorganism biofilm reactor"),
        "sidestream_pna":  ("Sidestream PN/A",               "Partial nitritation / anammox sidestream"),
    }
    # Try to get tech codes from scenario_names via decision data
    _dec = getattr(report, "decision", None)
    if _dec:
        _all_techs = list(getattr(_dec, "all_scenarios", {}).keys())
    else:
        _all_techs = []

    for name in (report.scenario_names or []):
        # Try to match name to a tech code
        tech_code = None
        tech_label = "—"
        key_char   = "—"
        # Check decision engine knowledge base
        for tc, (tl, kc) in _scenario_chars.items():
            if name.lower() in tl.lower() or tl.lower() in name.lower():
                tech_code  = tc
                tech_label = tl
                key_char   = kc
                break
        # Fallback: match by known name patterns
        if not tech_code:
            nl = name.lower()
            if "nereda" in nl or "granular" in nl or "ags" in nl:
                tech_label = "Aerobic Granular Sludge"; key_char = "SBR-based granular sludge, no clarifiers"
            elif "mabr" in nl:
                tech_label = "MABR + BNR"; key_char = "Membrane-aerated biofilm in BNR aerobic zone"
            elif "mbr" in nl and "bnr" in nl:
                tech_label = "BNR + MBR Hybrid"; key_char = "BNR basins with membrane separation"
            elif "mbbr" in nl or "ifas" in nl:
                tech_label = "IFAS / MBBR"; key_char = "Integrated fixed-film activated sludge"
            elif "bnr" in nl or "base" in nl:
                tech_label = "BNR (Activated Sludge)"; key_char = "Conventional BNR + secondary clarifiers"
        scen_rows.append([name, tech_label, key_char])
    if len(scen_rows) > 1:
        t = Table(scen_rows, colWidths=[W*0.35, W*0.35, W*0.30], repeatRows=1)
        t.setStyle(_pdf_tbl_style(colours))
        story.append(t)
        story.append(Spacer(1, 6))

    # ── Multi-criteria comparison ──────────────────────────────────────────
    if report.comparison_table:
        story.append(Paragraph("Multi-Criteria Comparison", styles["h1"]))
        story.append(_pdf_hr(colours))
        t = _render_list_table(report.comparison_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(Spacer(1, 4))

    # ── Key metrics KPI row ────────────────────────────────────────────────
    if report.cost_table:
        story.append(Spacer(1, 4))
        story.append(Paragraph("Cost Summary", styles["h1"]))
        story.append(_pdf_hr(colours))
        t = _render_dict_table(report.cost_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(Paragraph(
                "All CAPEX figures are concept-level estimates (±40%) in AUD 2024. "
                "Lifecycle cost = OPEX + CAPEX annualised over 30 years at 7% discount rate.",
                styles["disc"]))

    # ── Carbon & Energy ────────────────────────────────────────────────────
    if report.carbon_table:
        story.append(Spacer(1, 6))
        story.append(Paragraph("Carbon & Energy Summary", styles["h1"]))
        story.append(_pdf_hr(colours))
        t = _render_dict_table(report.carbon_table, styles, colours, W)
        if t:
            story.append(t)

    # ── Key conclusion ─────────────────────────────────────────────────────
    story.append(Spacer(1, 8))
    story.append(Paragraph("Conclusions", styles["h1"]))
    story.append(_pdf_hr(colours))
    # Prefer scoring_result preferred over is_preferred flag (scoring runs after exec summary build)
    _sr = getattr(report, "scoring_result", None)
    _exec_preferred = (
        _sr.preferred.scenario_name if _sr and _sr.preferred
        else report.preferred_scenario
    )
    if _exec_preferred:
        story.append(Paragraph(
            f"Based on the multi-criteria assessment, <b>{_exec_preferred}</b> "
            f"is the preferred option. Detailed engineering assessment "
            f"is recommended before proceeding to procurement.",
            styles["body"]))
    else:
        # Check decision engine result
        _dec2 = getattr(report, "decision", None)
        _dec_preferred = getattr(_dec2, "recommended_label", None) if _dec2 else None
        if _dec_preferred:
            story.append(Paragraph(
                f"Based on the multi-criteria assessment, <b>{_dec_preferred}</b> "
                f"is the preferred option under the compliance and cost framework. "
                f"Detailed engineering assessment is recommended before procurement.",
                styles["body"]))
        else:
            story.append(Paragraph(
                "All scenarios have been evaluated on a consistent basis. "
                "No single preferred option could be nominated — selection should "
                "be based on site-specific constraints, effluent standards, and "
                "stakeholder priorities.",
                styles["body"]))

    # ── Disclaimer ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 12))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        "DISCLAIMER: Concept-level estimates (±40%). Not for procurement, funding approval, "
        "or investment decisions without site-specific engineering by a qualified engineer. "
        "Prepared by ph2o consulting using the Water Utility Planning Platform.",
        styles["disc"]))

    doc.build(story, onFirstPage=hf, onLaterPages=hf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# COMPREHENSIVE REPORT PDF
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_comprehensive(report: ReportObject) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.platypus.doctemplate import BaseDocTemplate, PageTemplate
    from reportlab.platypus import Frame
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    styles, colours = _pdf_styles()
    buf = io.BytesIO()
    W = A4[0] - 36*mm
    H = A4[1]
    L, R, T, B = 18*mm, 18*mm, 22*mm, 22*mm

    # ── TOC-aware doc template ────────────────────────────────────────────
    # Add h1_toc style — h1 that registers in the TOC
    rl_blue = colours["blue"]
    rl_grey = colours["grey"]
    styles["h1_toc"] = ParagraphStyle(
        "rpt_h1_toc", parent=styles["h1"],
        fontSize=13, fontName="Helvetica-Bold", textColor=rl_blue,
        spaceBefore=14, spaceAfter=4,
    )

    class _TocDoc(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                style = flowable.style.name
                if style == "rpt_h1_toc":
                    txt = flowable.getPlainText()
                    self.notify("TOCEntry", (0, txt, self.page))

    doc = _TocDoc(buf, pagesize=A4,
                  leftMargin=L, rightMargin=R, topMargin=T, bottomMargin=B)

    frame = Frame(L, B, A4[0]-L-R, A4[1]-T-B, id="main")

    def hf(canvas, doc):
        _pdf_header_footer(canvas, doc, report)

    pt = PageTemplate(id="main", frames=[frame], onPage=hf)
    doc.addPageTemplates([pt])

    def H1(text):
        """Heading 1 that registers in TOC."""
        return Paragraph(text, styles["h1_toc"])

    story = []

    # ── Cover page ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 20*mm))
    story.append(Paragraph(
        report.project_name or "Treatment Technology Comparison", styles["title"]))
    story.append(Paragraph(
        report.plant_name or "Wastewater Treatment Facility", styles["subtitle"]))
    story.append(Spacer(1, 4))
    meta_rows = [
        ["Report Type:", "Concept-Stage Treatment Technology Comparison"],
        ["Domain:", report.domain or "Wastewater Treatment"],
        ["Scenarios:", str(len(report.scenario_names))],
        ["Prepared by:", report.prepared_by or "ph2o consulting"],
        ["Reviewed by:", report.reviewed_by or "—"],
        ["Date:", datetime.now().strftime("%d %B %Y")],
        ["Status:", "CONCEPT — ±40% ESTIMATE"],
    ]
    t = Table(meta_rows, colWidths=[W*0.30, W*0.70])
    t.setStyle(TableStyle([
        ("FONTNAME",  (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE",  (0,0), (-1,-1), 9),
        ("FONTNAME",  (0,0), (0,-1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0), (0,-1), colours["blue"]),
        ("TOPPADDING",    (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("LINEBELOW",     (0,-1), (-1,-1), 0.5, colours["blue"]),
    ]))
    from reportlab.platypus import TableStyle
    story.append(t)
    story.append(Spacer(1, 8*mm))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        "⚠ CONCEPT LEVEL ESTIMATE — ±40% — FOR COMPARATIVE PURPOSES ONLY. "
        "NOT FOR PROCUREMENT, FUNDING APPROVAL, OR INVESTMENT DECISIONS.",
        styles["disc"]))
    story.append(PageBreak())

    # ── Table of Contents ──────────────────────────────────────────────────
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOCHeading1", fontName="Helvetica", fontSize=9,
                       leftIndent=0, firstLineIndent=0, spaceBefore=4,
                       leading=14, textColor=colors.HexColor(f"#{GREY_HEX}")),
    ]
    story.append(P("Contents", styles["title"]))
    story.append(_pdf_hr(colours))
    story.append(toc)
    story.append(PageBreak())

    # ── Decision Summary Box ───────────────────────────────────────────────
    ds = getattr(report, "decision_summary", None)
    # Override decision_summary with scoring_result when available — scoring engine
    # is the single source of truth for the recommendation per the platform spec.
    _sr_box = getattr(report, "scoring_result", None)
    if _sr_box and _sr_box.preferred and ds:
        # Update preferred/runner-up from scoring engine
        ds = dict(ds)  # copy to avoid mutating the original
        ds["preferred"] = _sr_box.preferred.scenario_name
        ds["runner_up"] = _sr_box.runner_up.scenario_name if _sr_box.runner_up else ds.get("runner_up")
        # Update key driver to reflect scoring engine recommendation
        sr_pref = _sr_box.preferred
        sr_ru   = _sr_box.runner_up
        if sr_ru and sr_pref.criterion_scores:
            # Use LCC advantage as driver if it exists
            lcc_pref = sr_pref.criterion_scores.get("lcc")
            lcc_ru   = sr_ru.criterion_scores.get("lcc")
            if lcc_pref and lcc_ru:
                diff    = lcc_ru.raw_value - lcc_pref.raw_value  # positive = pref is cheaper
                op_pref = sr_pref.criterion_scores.get("operational_risk")
                op_ru   = sr_ru.criterion_scores.get("operational_risk")
                risk_note = ""
                if op_pref and op_ru:
                    risk_diff = int(op_ru.raw_value - op_pref.raw_value)  # positive = pref has lower risk
                    if abs(risk_diff) >= 3:
                        risk_note = f", {abs(risk_diff)} points lower risk" if risk_diff > 0 else f", {abs(risk_diff)} points higher risk"
                if diff > 0:
                    ds["driver"] = (
                        f"{sr_pref.scenario_name} saves ${diff:.0f}k/yr lifecycle cost "
                        f"vs {sr_ru.scenario_name}{risk_note}"
                    )
                else:
                    ds["driver"] = (
                        f"{sr_pref.scenario_name} costs ${abs(diff):.0f}k/yr more than "
                        f"{sr_ru.scenario_name} but scores higher on risk and maturity{risk_note}"
                    )
    if ds and ds.get("preferred"):
        from reportlab.platypus import Table as _DST, TableStyle as _DSTS
        from reportlab.lib import colors as _rlc

        rl_blue  = colours["blue"]
        rl_lt    = colours["lt"]

        # Header row
        hdr = [Paragraph("<b>🔷  DECISION SUMMARY</b>", styles["h1"])]
        ds_rows = [
            [Paragraph("<b>Preferred option</b>",  styles["h2"]),
             Paragraph(rl_safe(ds["preferred"]),    styles["body"])],
        ]
        if ds.get("runner_up"):
            ds_rows.append([
                Paragraph("<b>Runner-up</b>",           styles["h2"]),
                Paragraph(rl_safe(ds["runner_up"]),     styles["body"]),
            ])
        ds_rows.append([
            Paragraph("<b>Key driver</b>",              styles["h2"]),
            Paragraph(rl_safe(ds["driver"]),            styles["body"]),
        ])
        ds_rows.append([
            Paragraph("<b>Key trade-off</b>",           styles["h2"]),
            Paragraph(rl_safe(ds["trade_off"]),         styles["body"]),
        ])
        ds_rows.append([
            Paragraph("<b>Selection basis</b>",         styles["h2"]),
            Paragraph(rl_safe(ds["basis"]),             styles["body"]),
        ])
        ds_rows.append([
            Paragraph("<b>Confidence</b>",              styles["h2"]),
            Paragraph(rl_safe(ds.get("confidence","Moderate")), styles["body"]),
        ])
        if ds.get("runner_up_note"):
            ds_rows.append([
                Paragraph("<b>Runner-up preferred when</b>", styles["h2"]),
                Paragraph(rl_safe(ds["runner_up_note"]),     styles["body"]),
            ])

        col_w = [W * 0.28, W * 0.72]
        ds_tbl = _DST(ds_rows, colWidths=col_w)
        ds_tbl.setStyle(_DSTS([
            ("FONTNAME",      (0,0), (-1,-1), "Helvetica"),
            ("FONTSIZE",      (0,0), (-1,-1), 9),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 8),
            ("RIGHTPADDING",  (0,0), (-1,-1), 8),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS",(0,0), (-1,-1),
             [_rlc.HexColor("#EBF3FB"), _rlc.white]),
            ("BOX",           (0,0), (-1,-1), 1.5, rl_blue),
            ("LINEBELOW",     (0,0), (-1,-1), 0.25, _rlc.HexColor("#CCCCCC")),
            ("FONTNAME",      (0,0), (0,-1), "Helvetica-Bold"),
            ("TEXTCOLOR",     (0,0), (0,-1), rl_blue),
        ]))

        # Blue header bar
        # Header: ParagraphStyle with white text — table TEXTCOLOR doesn't override Paragraph colour
        _hdr_style = ParagraphStyle("ds_hdr", parent=styles["h2"],
                                    textColor=_rlc.white, fontName="Helvetica-Bold", fontSize=10)
        hdr_tbl = _DST([[Paragraph(
            "DECISION SUMMARY — " + rl_safe(ds["preferred"]) + " RECOMMENDED",
            _hdr_style)]],
            colWidths=[W])
        hdr_tbl.setStyle(_DSTS([
            ("BACKGROUND",    (0,0), (-1,-1), rl_blue),
            ("TOPPADDING",    (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
        ]))
        story.append(hdr_tbl)
        story.append(ds_tbl)
        story.append(Spacer(1, 12))

    # ── 1. Introduction ────────────────────────────────────────────────────
    story.append(H1("1. Introduction"))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        f"This report presents the results of a concept-stage treatment technology comparison "
        f"study for <b>{report.plant_name or 'the facility'}</b>. The study was undertaken using "
        f"the ph2o consulting Water Utility Planning Platform, which applies first-principles "
        f"engineering calculations based on Metcalf &amp; Eddy (5th Edition), WEF Manual of "
        f"Practice MOP 32 and MOP 35, and published technology-specific references.",
        styles["body"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "The purpose of this study is to provide a structured, consistent basis for comparing "
        "treatment technology options at concept stage. All cost estimates are ±40% and are "
        "not suitable for procurement, detailed feasibility, or funding approval without "
        "site-specific investigation by a qualified engineer.",
        styles["body"]))
    story.append(Spacer(1, 8))

    # ── 2. Study Scope ─────────────────────────────────────────────────────
    story.append(H1("2. Study Scope and Methodology"))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        f"A total of <b>{len(report.scenario_names)} treatment scenarios</b> were evaluated:",
        styles["body"]))
    for name in (report.scenario_names or []):
        story.append(P(f"• {name}", styles["bullet"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "Each scenario was evaluated on the following criteria:",
        styles["body"]))
    criteria = [
        "Capital cost (CAPEX) — civil, mechanical, electrical, and membrane works",
        "Operating cost (OPEX) — electricity, chemicals, sludge disposal, maintenance",
        "Lifecycle cost — CAPEX annualised at 7% over 30 years plus OPEX",
        "Specific energy consumption (kWh/ML treated)",
        "Carbon footprint — Scope 1 (process emissions) and Scope 2 (grid electricity)",
        "Effluent quality — BOD, TSS, TN, NH4, TP",
        "Biosolids production and sludge yield",
        "Technology risk — maturity, implementation, operational complexity, regulatory",
    ]
    for c in criteria:
        story.append(P(f"• {c}", styles["bullet"]))
    story.append(Spacer(1, 8))

    # ── 3. Multi-Criteria Comparison ───────────────────────────────────────
    if report.comparison_table:
        story.append(H1("3. Multi-Criteria Comparison"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 1 presents the key comparative metrics across all evaluated scenarios.",
            styles["body"]))
        t = _render_list_table(report.comparison_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(P("Table 1: Multi-criteria comparison summary", styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 4. Process Design Summary ─────────────────────────────────────────
    story.append(H1("4. Process Design Summary"))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        "Table 2 presents the key process design parameters for each evaluated scenario, "
        "derived from first-principles calculations using standard biological treatment "
        "design methods (Metcalf & Eddy 5th Edition).",
        styles["body"]))

    # Build Table 2 from scenario_design_data
    design_data = getattr(report, "scenario_design_data", {})
    if design_data:
        from reportlab.platypus import Table as _DT, TableStyle as _DTS
        DESIGN_PARAMS = [
            ("Bioreactor volume",   "reactor_volume_m3",          "{:,.0f} m\u00b3"),
            ("Process footprint",   "footprint_m2",               "{:,.0f} m\u00b2"),
            ("MLSS",                "mlss_granular_mg_l",         "{:,.0f} mg/L"),
            ("O\u2082 demand",      "o2_demand_kg_day",           "{:,.0f} kg/day"),
            ("Sludge production",   "sludge_production_kgds_day", "{:,.0f} kgDS/day"),
            ("Specific energy",     "energy_intensity_kwh_kl",    "{:.0f} kWh/ML"),
        ]
        scen_names = list(design_data.keys())
        # Header row
        tbl_rows = [[Paragraph("Parameter", styles["h2"])] +
                    [Paragraph(n, styles["h2"]) for n in scen_names]]
        for label, key, fmt in DESIGN_PARAMS:
            row_vals = []
            for scen_name in scen_names:
                sd = design_data[scen_name]
                tech_code = sd.get("tech_sequence", [None])[0]
                perf = sd.get("tech_performance", {}).get(tech_code, {}) if tech_code else {}
                val = perf.get(key)
                # Special case: energy_intensity is in kWh/kL, display as kWh/ML
                if key == "energy_intensity_kwh_kl" and val is not None:
                    val = val * 1000
                if val is not None and val > 0:
                    try:
                        cell = fmt.format(val)
                    except Exception:
                        cell = str(val)
                else:
                    cell = None  # Use None to detect all-blank rows
                row_vals.append(cell)

            # Skip MLSS row if no scenario has it (non-AGS technologies)
            if key == "mlss_granular_mg_l" and all(v is None for v in row_vals):
                continue

            row = [Paragraph(rl_safe(label), styles["body"])]
            for cell in row_vals:
                row.append(Paragraph(rl_safe(cell or "—"), styles["body"]))
            tbl_rows.append(row)

        # Effluent quality row
        eff_row = [Paragraph("Effluent quality", styles["body"])]
        for scen_name in scen_names:
            sd = design_data[scen_name]
            tech_code = sd.get("tech_sequence", [None])[0]
            perf = sd.get("tech_performance", {}).get(tech_code, {}) if tech_code else {}
            parts = []
            for ek, elbl in [("effluent_bod_mg_l","BOD"),("effluent_tss_mg_l","TSS"),
                              ("effluent_tn_mg_l","TN"),("effluent_nh4_mg_l","NH4")]:
                v = perf.get(ek)
                if v is not None:
                    parts.append(f"{elbl} {v:.0f}")
            eff_row.append(Paragraph(rl_safe("  ".join(parts) + " mg/L") if parts else "—",
                                     styles["body"]))
        tbl_rows.append(eff_row)

        n_cols = 1 + len(scen_names)
        # Wider first column and more padding to prevent cell text concatenation
        col_w = [W * 0.30] + [W * 0.70 / len(scen_names)] * len(scen_names)
        t2 = _DT(tbl_rows, colWidths=col_w, repeatRows=1)
        t2.setStyle(_pdf_tbl_style(colours))
        story.append(Spacer(1, 4))
        story.append(t2)
        story.append(P("Table 2: Process design summary", styles["caption"]))

        # ── Engineering notes ──────────────────────────────────────────────
        # O2 vs energy explanation — always include when both AGS and BNR present
        has_ags = any("granular" in (design_data.get(n, {}).get("tech_sequence") or [""])[0]
                      for n in scen_names)
        has_bnr = any("bnr" == (design_data.get(n, {}).get("tech_sequence") or [""])[0]
                      for n in scen_names)

        if has_ags and has_bnr:
            story.append(P(
                "<b>Note — O₂ demand vs total energy:</b> Aerobic Granular Sludge (AGS/Nereda) "
                "shows higher O₂ demand than conventional BNR despite lower total plant energy. "
                "This is thermodynamically correct: AGS operates at lower sludge yield (y_obs ≈ 0.22 "
                "vs BNR ≈ 0.33 kgVSS/kgBOD), meaning more substrate is fully oxidised rather than "
                "incorporated into biomass — requiring more aeration oxygen per kg BOD removed. "
                "AGS wins on total plant energy because it eliminates secondary clarifiers, RAS "
                "recirculation, and MLR pumping, which together account for 30–40% of BNR plant energy.",
                styles["disc"]))
            story.append(Spacer(1, 4))

        # TN caveat — when BNR achieves TN=5 at low C:N ratio, flag the assumption
        for scen_name in scen_names:
            sd = design_data.get(scen_name, {})
            tech_code = (sd.get("tech_sequence") or [None])[0]
            if tech_code != "bnr":
                continue
            perf = sd.get("tech_performance", {}).get(tech_code, {})
            eff_tn = perf.get("effluent_tn_mg_l", 99)
            domain_inp = sd.get("domain_inputs", {})
            bod = domain_inp.get("influent_bod_mg_l", 250)
            tkn = domain_inp.get("influent_tkn_mg_l", 45) or 1
            cn = bod / tkn
            if eff_tn <= 5.5 and cn < 8.0:
                story.append(P(
                    f"<b>Note — BNR effluent TN assumption:</b> Achieving TN ≤ {eff_tn:.0f} mg/L "
                    f"in conventional BNR at BOD/TKN = {cn:.1f} requires adequate carbon for "
                    "denitrification. At this C:N ratio, supplemental carbon (methanol or "
                    "VFA from primary sludge fermentation), tight SRT control, and favourable "
                    "temperature conditions are assumed. Without these, BNR effluent TN may be "
                    "10–15 mg/L. Confirm carbon availability during site investigation.",
                    styles["disc"]))
                story.append(Spacer(1, 4))

    story.append(Spacer(1, 6))

    # ── PFD section ────────────────────────────────────────────────────────
    story.append(Spacer(1, 6))
    story.append(H1("5. Process Flow Diagrams"))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        "The following schematic process flow diagrams illustrate the treatment train "
        "configuration for each evaluated scenario. Diagrams show principal unit processes, "
        "key recycle streams (RAS, MLR, WAS), aeration systems, and sludge handling. "
        "Diagrams are schematic only and not to scale.",
        styles["body"]))
    story.append(Spacer(1, 4))

    try:
        from core.reporting.pfd_generator import get_pfd
        from reportlab.lib.units import mm
        pfd_w = W
        pfd_h = 155

        for sec in (report.sections or []):
            pass  # placeholder

        # Draw PFD for each scenario if tech code available
        design_data = getattr(report, "scenario_design_data", {})

        for scen_idx, scen_name in enumerate(report.scenario_names or []):
            sd = design_data.get(scen_name, {})
            tech_seq = sd.get("tech_sequence", [])
            tp_all   = sd.get("tech_performance", {})

            # Primary biological tech is first in sequence
            tech_code = tech_seq[0] if tech_seq else None
            perf      = tp_all.get(tech_code, {}) if tech_code else {}

            if tech_code:
                story.append(P(f"Scenario: {scen_name}", styles["h2"]))
                story.append(Spacer(1, 2))

                # ── PFD drawing ───────────────────────────────────────────
                pfd = get_pfd(tech_code, perf, width=pfd_w, height=pfd_h)
                story.append(pfd)
                story.append(Paragraph(
                    f"Figure {scen_idx + 1} — Schematic PFD: {scen_name}. "
                    "Streams: ── Process flow  - - - Recycle (orange)  → Effluent (green)  → Sludge (red). Schematic only.",
                    styles["caption"]))
                story.append(Spacer(1, 10))
    except Exception:
        story.append(Paragraph(
            "Process flow diagrams not available for this report.", styles["body"]))

    # ── Re-number remaining sections ────────────────────────────────────────
    # ── 6. Capital Cost ─────────────────────────────────────────────────────
    if report.cost_table:
        story.append(H1("6. Capital and Lifecycle Cost Assessment"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 3 summarises the capital cost, operating cost, and lifecycle cost for each "
            "scenario. All costs are in AUD 2024 and are concept-level estimates (±40%).",
            styles["body"]))
        t = _render_dict_table(report.cost_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(P("Table 3: Cost summary (AUD 2024, concept-level ±40%)", styles["caption"]))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            "CAPEX estimates cover bioreactor civil works, mechanical and electrical equipment, "
            "membranes (where applicable), secondary clarifiers, and instrumentation. They exclude "
            "land, site preparation, headworks, sludge treatment train, buildings, electrical "
            "reticulation, owner's costs, and escalation. A 20% design contingency and 12% "
            "contractor margin are included in the unit rates.",
            styles["body"]))
        # CAPEX + LCC side-by-side charts
        chart_row = []
        if report.charts.get("capex_comparison"):
            ch = _chart_capex(report.charts["capex_comparison"], W * 0.48, height=150)
            if ch:
                chart_row.append(ch)
        # LCC chart — build from cost_table data
        lcc_data = {"x": [], "y": []}
        if report.cost_table:
            for row in report.cost_table.get("rows", []):
                lcc_data["x"].append(row.get("Scenario", ""))
                try:
                    lcc_data["y"].append(float(str(row.get("Lifecycle Cost ($/yr)", "0")).replace(",", "")) / 1e3)
                except (ValueError, TypeError):
                    lcc_data["y"].append(0)
        if lcc_data["x"] and any(v > 0 for v in lcc_data["y"]):
            lcc_data["ylabel"] = "LCC (k$/yr)"
            ch_lcc = _chart_capex(lcc_data, W * 0.48, height=150, label_fmt="${:.0f}k/yr")
            if ch_lcc:
                chart_row.append(ch_lcc)
        if chart_row:
            from reportlab.platypus import Table as _T, TableStyle as _TS
            if len(chart_row) == 2:
                ct = _T([[chart_row[0], chart_row[1]]], colWidths=[W*0.49, W*0.49])
                ct.setStyle(_TS([("VALIGN",(0,0),(-1,-1),"TOP"),
                                  ("LEFTPADDING",(0,0),(-1,-1),0),
                                  ("RIGHTPADDING",(0,0),(-1,-1),4)]))
                story.append(Spacer(1, 4))
                story.append(ct)
                story.append(P(
                    "Figure A — CAPEX ($M) and Lifecycle Cost (k$/yr) comparison by scenario (AUD 2024)",
                    styles["caption"]))
            else:
                story.append(Spacer(1, 4))
                story.append(chart_row[0])
                story.append(P("Figure A — CAPEX comparison by scenario ($M, AUD 2024)", styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 6a. OPEX Breakdown ────────────────────────────────────────────────
    if getattr(report, "opex_breakdown_table", None):
        story.append(H1("6a. OPEX Cost Drivers"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 3a breaks down annual operating cost by category. This identifies the primary "
            "cost drivers and explains the OPEX differential between scenarios. "
            "Energy and sludge disposal are typically the two largest OPEX components for "
            "biological treatment processes.",
            styles["body"]))
        t = _render_dict_table(report.opex_breakdown_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(Paragraph(
                "Table 3a: OPEX breakdown by category (AUD 2024/yr). "
                "Percentages shown as proportion of total annual OPEX.",
                styles["caption"]))
        # OPEX stacked chart — build from opex_breakdown chart data
        # Use shortened category labels so legend fits without truncation
        opex_cd = report.charts.get("opex_breakdown", {})
        if opex_cd and opex_cd.get("scenarios") and opex_cd.get("categories"):
            _label_short = {
                "Operator & maintenance labour": "Labour",
                "Electricity — aeration":        "Electricity (aeration)",
                "Electricity — mixing & decant": "Electricity (mix/decant)",
                "Sludge disposal":               "Sludge disposal",
            }
            cats = opex_cd["categories"]
            short_cats = [_label_short.get(c, c[:20]) for c in cats]
            scen_data = {
                scen: {short_cats[i]: opex_cd["data"].get(cat, [0]*len(opex_cd["scenarios"]))[si]
                       for i, cat in enumerate(cats)}
                for si, scen in enumerate(opex_cd["scenarios"])
            }
            ch = _chart_stacked_h(short_cats, scen_data, W,
                                   height=max(130, len(opex_cd["scenarios"]) * 60 + 50))
            if ch:
                story.append(Spacer(1, 4))
                story.append(ch)
                story.append(P("Figure B — Annual OPEX breakdown by category (k$/yr)", styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 6b. Specific Performance Metrics ──────────────────────────────────
    if getattr(report, "specific_metrics_table", None):
        story.append(H1("6b. Specific Performance Metrics"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 3b presents normalised performance metrics to enable direct comparison "
            "across scenarios and benchmarking against industry references. "
            "Specific footprint (m2/MLD) and specific sludge (kgDS/ML) are standard "
            "planning metrics used in Australian utility capital planning.",
            styles["body"]))
        t = _render_dict_table(report.specific_metrics_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(Paragraph(
                "Table 3b: Specific performance metrics. Carbon intensity in kgCO2e/kL "
                "enables direct comparison with water supply carbon benchmarks.",
                styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 7. Energy & Carbon ─────────────────────────────────────────────────
    if report.carbon_table:
        story.append(H1("7. Energy and Carbon Footprint"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 4 summarises the carbon emissions for each scenario. Scope 1 emissions "
            "include N2O from biological nitrogen removal and CH4 from sludge handling. "
            "Scope 2 emissions are calculated from grid electricity consumption using the "
            "applicable emission factor. Avoided emissions reflect CHP electricity generation "
            "where anaerobic digestion is included.",
            styles["body"]))
        t = _render_dict_table(report.carbon_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(P("Table 4: Carbon and energy summary", styles["caption"]))
        # Carbon chart — Scope 1 vs Scope 2 stacked
        carbon_cd = report.charts.get("carbon_comparison", {})
        if carbon_cd and carbon_cd.get("scenarios"):
            scens = carbon_cd["scenarios"]
            scen_data = {
                scen: {
                    "Scope 1 (N2O/CH4)": carbon_cd["scope_1"][si],
                    "Scope 2 (Electricity)": carbon_cd["scope_2"][si],
                }
                for si, scen in enumerate(scens)
            }
            labels = ["Scope 1 (N2O/CH4)", "Scope 2 (Electricity)"]
            ch = _chart_stacked_h(labels, scen_data, W,
                                   height=max(120, len(scens) * 55 + 40),
                                   unit="tCO2e")
            if ch:
                story.append(Spacer(1, 4))
                story.append(ch)
                story.append(Paragraph(
                    "Figure C — Carbon footprint by scenario and scope (tCO2e/yr)",
                    styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 8. Risk Assessment ─────────────────────────────────────────────────
    if report.risk_table:
        story.append(H1("8. Risk Assessment"))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "Table 5 presents the technology risk assessment across technical, implementation, "
            "operational, and regulatory dimensions. Risk scores are indicative and based on "
            "technology maturity, reference plant availability, and operational complexity.",
            styles["body"]))
        t = _render_dict_table(report.risk_table, styles, colours, W)
        if t:
            story.append(t)
            story.append(P("Table 5: Technology risk assessment", styles["caption"]))
        story.append(Spacer(1, 8))

    # ── 7. Conclusions ─────────────────────────────────────────────────────
    story.append(H1("9. Conclusions and Recommendations"))
    story.append(_pdf_hr(colours))

    # ── Identify compliance status from comparison table ──────────────────
    non_compliant = []
    compliant     = []
    tn_vals       = {}
    if report.comparison_table:
        # Scan for scenarios where effluent TN doesn't meet 5.0 mg/L target
        for row in report.comparison_table:
            if row.get("Criterion") == "Effluent TN (mg/L)":
                for scen in report.scenario_names:
                    try:
                        val = float(str(row.get(scen, "99")).replace("—","99"))
                        tn_vals[scen] = val
                    except (ValueError, TypeError):
                        pass

    # ── Non-compliance block ───────────────────────────────────────────────
    if tn_vals:
        for scen, val in tn_vals.items():
            if val > 5.5:   # 10% tolerance on 5.0 mg/L target
                non_compliant.append((scen, val))
            else:
                compliant.append(scen)

        if non_compliant:
            story.append(P(
                "<b>⚠ COMPLIANCE NOTE — READ BEFORE PROCEEDING</b>",
                styles["h2"]))
            for scen, tn in non_compliant:
                story.append(P(
                    f"<b>{scen}</b> does <b>not meet the TN licence target</b> as modelled "
                    f"(effluent TN = {tn:.1f} mg/L vs target 5.0 mg/L). "
                    "This scenario is non-compliant as a standalone option and "
                    "<b>must not be recommended for procurement without engineering intervention</b> "
                    "(supplemental carbon dosing, SRT extension, or technology modification). "
                    "See alternative pathways in Section 3.",
                    styles["body"]))
            story.append(Spacer(1, 4))

    # ── Preferred option ──────────────────────────────────────────────────
    if report.preferred_scenario:
        story.append(P(
            f"<b>Recommended option: {report.preferred_scenario}</b>",
            styles["h2"]))
        story.append(P(
            f"Based on the multi-criteria assessment, <b>{report.preferred_scenario}</b> "
            "is the preferred option on lifecycle cost. "
            + (f"Note: {', '.join(s for s,_ in non_compliant)} is excluded from recommendation "
               "due to compliance failure. " if non_compliant else ""),
            styles["body"]))
    elif compliant:
        # Find the lowest LCC scenario among compliant options
        lcc_winner = None
        lcc_min = float("inf")
        lcc_delta_k = None
        if report.cost_table:
            for row in report.cost_table.get("rows", []):
                scen = row.get("Scenario", "")
                if scen not in compliant:
                    continue
                try:
                    lcc = float(str(row.get("Lifecycle Cost ($/yr)", "0")).replace(",", ""))
                    if lcc < lcc_min:
                        lcc_min = lcc
                        lcc_winner = scen
                except (ValueError, TypeError):
                    pass
            # Calculate delta vs second-best
            second_lcc = float("inf")
            for row in report.cost_table.get("rows", []):
                scen = row.get("Scenario", "")
                if scen not in compliant or scen == lcc_winner:
                    continue
                try:
                    lcc = float(str(row.get("Lifecycle Cost ($/yr)", "0")).replace(",", ""))
                    second_lcc = min(second_lcc, lcc)
                except (ValueError, TypeError):
                    pass
            if second_lcc < float("inf") and lcc_winner:
                lcc_delta_k = round((second_lcc - lcc_min) / 1e3)

        story.append(P(
            f"<b>Compliant options: {', '.join(compliant)}</b>",
            styles["h2"]))
        if lcc_winner and lcc_delta_k and lcc_delta_k > 0:
            story.append(P(
                f"Both scenarios achieve the effluent quality targets as modelled. "
                f"On lifecycle cost, <b>{lcc_winner}</b> is the preferred option, "
                f"saving approximately <b>${lcc_delta_k:,}k/yr</b> over the 30-year analysis period "
                f"compared to the alternative. "
                f"Final selection should confirm this advantage holds under site-specific conditions.",
                styles["body"]))
        else:
            story.append(P(
                "All compliant scenarios achieve the effluent quality targets as modelled. "
                "Final selection should be based on lifecycle cost comparison and site-specific investigation.",
                styles["body"]))
    else:
        story.append(P(
            "<b>No scenario achieves full compliance as modelled.</b> "
            "Engineering interventions are required before any option can be recommended for procurement.",
            styles["body"]))

    story.append(Spacer(1, 4))

    # ── Selection guidance ────────────────────────────────────────────────
    story.append(P("Final selection should consider:", styles["body"]))
    recs = [
        "Site-specific constraints: footprint, existing infrastructure, and staging requirements",
        "Regulatory licence conditions: confirm effluent TN, TP, and pathogen limits with the relevant EPA",
        "Utility priorities: weight energy, carbon, capital cost, and risk according to organisational strategy",
        "Procurement strategy: D&C suitability, local contractor capability, and technology vendor engagement",
        "Detailed geotechnical, hydraulic, and site investigation before design or cost confirmation",
    ]
    for rec in recs:
        story.append(P(f"• {rec}", styles["bullet"]))
    story.append(Spacer(1, 4))
    story.append(P(
        "A detailed feasibility study with site-specific cost estimation (±15–20%) "
        "is recommended before proceeding to detailed design or procurement. "
        "All CAPEX and OPEX figures in this report are concept-level estimates (±40%) "
        "suitable for comparative ranking only.",
        styles["body"]))
    story.append(Spacer(1, 8))

    # ── Appendix B — Decision Scoring ────────────────────────────────────
    sr = getattr(report, "scoring_result", None)
    if sr and sr.preferred:
        from reportlab.platypus import Table as _ST, TableStyle as _STS
        story.append(PageBreak())
        story.append(P("Appendix B — Decision Scoring", styles["h1"]))
        story.append(_pdf_hr(colours))
        story.append(P(
            f"The following weighted multi-criteria decision analysis was performed using the "
            f"<b>{sr.weight_profile}</b> weight profile. Scores are normalised 0–100 per criterion "
            "across the compared scenarios. Higher scores indicate better performance. "
            "Compliance is a hard gate — non-compliant options are excluded from recommendation.",
            styles["body"]))
        story.append(Spacer(1, 6))

        # Decision hierarchy explanation
        story.append(P("<b>Decision Logic</b>", styles["h2"]))
        excl_names = ", ".join(e.scenario_name for e in sr.excluded) if sr.excluded else "none"
        pref_name  = sr.preferred.scenario_name if sr.preferred else "—"
        pref_score = f"{sr.preferred.total_score:.0f}/100" if sr.preferred else "—"
        close_txt  = (f"Close decision (gap {sr.preferred.total_score - sr.runner_up.total_score:.1f} pts) "
                      "— validate before commitment"
                      if sr.close_decision and sr.runner_up
                      else "Clear preference — proceed with validation")
        _hier_rows = [
            [P("<b>Step</b>", styles["body"]), P("<b>Rule</b>", styles["body"]),
             P("<b>Outcome</b>", styles["body"])],
            [P("1. Compliance gate", styles["body"]),
             P("Fail any effluent target → excluded from recommendation", styles["body"]),
             P(rl_safe(f"Excluded: {excl_names}"), styles["body"])],
            [P("2. Cost ranking", styles["body"]),
             P("Rank compliant options by lifecycle cost (lower = preferred)", styles["body"]),
             P(rl_safe(f"Lowest LCC compliant: {pref_name}"), styles["body"])],
            [P("3. Weighted scoring", styles["body"]),
             P("Apply cost, environmental, risk and maturity weights", styles["body"]),
             P(rl_safe(f"Preferred: {pref_name} ({pref_score})"), styles["body"])],
            [P("4. Close-decision", styles["body"]),
             P(rl_safe(f"Gap < {sr.close_decision_threshold:.0f} pts → flag, do not lock in"), styles["body"]),
             P(rl_safe(close_txt), styles["body"])],
        ]
        from reportlab.lib import colors as _rlc3
        _hier_tbl = _ST(_hier_rows, colWidths=[W*0.20, W*0.47, W*0.33], repeatRows=1)
        _hier_tbl.setStyle(_STS([
            ("BACKGROUND",    (0,0), (-1,0), colours["blue"]),
            ("TEXTCOLOR",     (0,0), (-1,0), _rlc3.white),
            ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1,-1), 7.5),
            ("TOPPADDING",    (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("LEFTPADDING",   (0,0), (-1,-1), 5),
            ("GRID",          (0,0), (-1,-1), 0.25, colours["lt"]),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [colours["lt"], _rlc3.white]),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ]))
        story.append(_hier_tbl)
        story.append(Spacer(1, 8))

        # Recommendation paragraph
        story.append(P(f"<b>Preferred option:</b> {sr.preferred.scenario_name} "
                       f"({sr.preferred.total_score:.0f}/100)", styles["body"]))
        if sr.runner_up:
            story.append(P(f"<b>Runner-up:</b> {sr.runner_up.scenario_name} "
                           f"({sr.runner_up.total_score:.0f}/100)", styles["body"]))
        story.append(P(f"<b>Recommendation:</b> {sr.recommendation}", styles["body"]))
        if sr.trade_off:
            story.append(P(f"<b>Trade-off:</b> {sr.trade_off}", styles["body"]))
        story.append(Spacer(1, 6))

        # Close decision note
        if sr.close_decision and sr.runner_up:
            gap = sr.preferred.total_score - sr.runner_up.total_score
            story.append(P(
                f"⚠ Close decision: scoring gap is {gap:.1f} points (threshold 5.0). "
                "Final selection should depend on site-specific conditions and risk appetite, "
                "not score alone.",
                styles["disc"]))
            story.append(Spacer(1, 4))

        # Weighted score table
        from core.decision.scoring_engine import CRITERION_LABELS, WEIGHT_PROFILES, WeightProfile
        crit_keys = list(CRITERION_LABELS.keys())
        crit_lbls = [CRITERION_LABELS[k] for k in crit_keys]
        weights_used = sr.weights

        # Build header
        col_w_s = [W * 0.22] + [W * 0.10] + [W * 0.68 / len(crit_keys)] * len(crit_keys)
        score_hdr = (
            [Paragraph("Scenario", styles["h2"]),
             Paragraph("Total", styles["h2"])] +
            [Paragraph(rl_safe(f"{lbl[:14]}\n{int(weights_used.get(k,0)*100)}%"),
                       styles["caption"])
             for k, lbl in zip(crit_keys, crit_lbls)]
        )
        score_rows = [score_hdr]
        for opt in sr.scored_options:
            prefix = "★ " if opt.rank == 1 else ("✗ " if not opt.is_eligible else "")
            row = [
                Paragraph(rl_safe(prefix + opt.scenario_name), styles["body"]),
                Paragraph(
                    f"<b>{opt.total_score:.0f}</b>" if opt.is_eligible else
                    f"<i>{opt.total_score:.0f}</i>",
                    styles["body"]),
            ]
            for k in crit_keys:
                cs = opt.criterion_scores.get(k)
                row.append(Paragraph(
                    rl_safe(f"{cs.normalised:.0f}" if cs else "—"), styles["body"]))
            score_rows.append(row)

        score_tbl = _ST(score_rows, colWidths=col_w_s, repeatRows=1)
        from reportlab.lib import colors as _rlc2
        score_tbl.setStyle(_STS([
            ("BACKGROUND",    (0,0), (-1,0), colours["blue"]),
            ("TEXTCOLOR",     (0,0), (-1,0), _rlc2.white),
            ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1,-1), 7),
            ("TOPPADDING",    (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ("LEFTPADDING",   (0,0), (-1,-1), 4),
            ("GRID",          (0,0), (-1,-1), 0.25, colours["lt"]),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [colours["lt"], _rlc2.white]),
        ]))
        story.append(score_tbl)
        story.append(P(
            "Table B1: Weighted decision scores by criterion. "
            "★ = preferred option. ✗ = excluded (non-compliant).",
            styles["caption"]))
        story.append(Spacer(1, 6))

        # Excluded options
        if sr.excluded:
            story.append(P("<b>Excluded options:</b>", styles["body"]))
            for opt in sr.excluded:
                story.append(P(
                    f"• <b>{opt.scenario_name}</b> — {opt.excluded_reason}",
                    styles["disc"]))
        story.append(Spacer(1, 4))
        story.append(P(
            f"Weight profile: {sr.weight_profile}. Scores are concept-level (±40%). "
            "Sensitivity test with alternative profiles before technology lock-in.",
            styles["disc"]))


        # ── Intervention Scenarios ────────────────────────────────────────
        from reportlab.lib import colors as _rlc4
        _ivs = getattr(report, "intervention_results", None) or []
        if _ivs:
            story.append(Spacer(1, 10))
            story.append(P("<b>Intervention Scenarios</b>", styles["h2"]))
            story.append(P(
                "Engineering interventions that can make non-compliant options viable. "
                "Costs are indicative (\u00b140%). Ferric chloride dosing achieves TP target.",
                styles["body"]))
            story.append(Spacer(1, 4))
            _int_rows = [[
                P("<b>Option</b>", styles["body"]),
                P("<b>Intervention</b>", styles["body"]),
                P("<b>+OPEX (k$/yr)</b>", styles["body"]),
                P("<b>+CAPEX ($M)</b>", styles["body"]),
                P("<b>LCC with intervention</b>", styles["body"]),
                P("<b>Compliant?</b>", styles["body"]),
            ]]
            for _iv in _ivs:
                _ms  = _iv.modified_scenario
                _lcc = f"${_ms.cost_result.lifecycle_cost_annual/1e3:.0f}k/yr" if _ms and _ms.cost_result else "—"
                _ach = "Yes" if _iv.achieves_compliance else "Partial"
                _int = _iv.intervention_label.replace(_iv.base_scenario_name + " + ", "")
                _int_rows.append([
                    P(rl_safe(_iv.base_scenario_name), styles["body"]),
                    P(rl_safe(_int[:60]), styles["body"]),
                    P(rl_safe(f"+${_iv.opex_delta_k_yr:.0f}"), styles["body"]),
                    P(rl_safe(f"+${_iv.capex_delta_m:.2f}"), styles["body"]),
                    P(rl_safe(_lcc), styles["body"]),
                    P(rl_safe(_ach), styles["body"]),
                ])
            _int_tbl = _ST(_int_rows, colWidths=[W*.17,W*.28,W*.13,W*.12,W*.17,W*.13], repeatRows=1)
            _int_tbl.setStyle(_STS([
                ("BACKGROUND",    (0,0),(-1,0), colours["blue"]),
                ("TEXTCOLOR",     (0,0),(-1,0), _rlc4.white),
                ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
                ("FONTSIZE",      (0,0),(-1,-1), 7.5),
                ("TOPPADDING",    (0,0),(-1,-1), 4),
                ("BOTTOMPADDING", (0,0),(-1,-1), 4),
                ("LEFTPADDING",   (0,0),(-1,-1), 5),
                ("GRID",          (0,0),(-1,-1), 0.25, colours["lt"]),
                ("ROWBACKGROUNDS",(0,1),(-1,-1), [colours["lt"], _rlc4.white]),
                ("VALIGN",        (0,0),(-1,-1), "TOP"),
            ]))
            story.append(_int_tbl)

        # ── Carbon Decision Pathway ───────────────────────────────────────
        _cp = getattr(report, "carbon_pathway_result", None)
        _sr = getattr(report, "scoring_result", None)
        if _cp and _cp.preferred:
            story.append(Spacer(1, 10))
            story.append(P("<b>Carbon Decision Pathway — Low-Carbon Profile</b>", styles["h2"]))
            story.append(P(
                "Re-ranking under the Low-carbon / future-focused weight profile "
                "(carbon intensity 25%, cost 25%, risk 20%). "
                "Relevant if the utility has a carbon reduction target or expects "
                "carbon pricing to increase materially.",
                styles["body"]))
            story.append(Spacer(1, 4))
            _bal = _sr.preferred.scenario_name if _sr and _sr.preferred else "—"
            _lc  = _cp.preferred.scenario_name
            story.append(P(rl_safe(
                f"Balanced profile preferred: <b>{_bal}</b>    |    "
                f"Low-carbon profile preferred: <b>{_lc}</b> ({_cp.preferred.total_score:.0f}/100)"
            ), styles["body"]))
            if _lc != _bal:
                story.append(Spacer(1, 3))
                story.append(P(rl_safe(
                    f"Note: under carbon prioritisation, {_lc} becomes preferred over {_bal}. "
                    "If carbon reduction is a strategic priority, evaluate both options "
                    "before technology selection."
                ), styles["body"]))

        # ── Platform QA ───────────────────────────────────────────────────
        _qa = getattr(report, "platform_qa_result", None)
        if _qa:
            story.append(Spacer(1, 10))
            story.append(P("<b>Platform QA</b>", styles["h2"]))
            _qa_status = "PASSED" if _qa.passed else "FAILED"
            story.append(P(rl_safe(f"Status: {_qa_status}   "
                f"Errors: {len(_qa.errors)}   Warnings: {len(_qa.warnings)}   Notes: {len(_qa.notes)}"),
                styles["body"]))
            for _e in _qa.errors:
                story.append(P(rl_safe(f"ERROR: {_e[:120]}"), styles["body"]))
            for _w in _qa.warnings:
                story.append(P(rl_safe(f"Warning: {_w[:120]}"), styles["body"]))

    # ── Appendix A — Key Assumptions ────────────────────────────────────────
    if report.assumptions_appendix:
        story.append(PageBreak())
        story.append(P("Appendix A — Key Assumptions", styles["h1"]))
        story.append(_pdf_hr(colours))
        story.append(Paragraph(
            "The following key assumptions were used in the analysis. User-entered "
            "site-specific values are noted where applied.",
            styles["body"]))
        story.append(Spacer(1, 4))

        # Clean up and format assumptions — skip internal/debugging entries
        skip_prefixes = ["capex_unit_costs", "pump_per_kw", "blower_per_kw",
                         "vsd_per_kw", "instrumentation", "design_contingency",
                         "contractor_margin", "client_oncosts", "return_sludge",
                         "gravity_thickener", "belt_thickener", "centrifuge",
                         "struvite", "fine_screen", "mbr_membrane_flat"]
        clean_rows = []
        for item in report.assumptions_appendix:
            param = str(item.get("Parameter", ""))
            # Skip raw unit cost keys — shown in screenshot as layout-breakers
            if any(param.startswith(p) for p in skip_prefixes):
                continue
            clean_rows.append({
                "Category":  item.get("Category", ""),
                "Parameter": _format_assumption_param(param),
                "Value":     str(item.get("Value", "")),
                "Override":  item.get("User Override", ""),
            })

        if clean_rows:
            hdrs = ["Category", "Parameter", "Value", "Override"]
            tbl_data = [[Paragraph(h, styles["h2"]) for h in hdrs]]
            for r in clean_rows:
                tbl_data.append([
                    Paragraph(r["Category"],  styles["body"]),
                    Paragraph(r["Parameter"], styles["body"]),
                    Paragraph(r["Value"],     styles["body"]),
                    Paragraph(r["Override"],  styles["body"]),
                ])
            t = Table(tbl_data,
                      colWidths=[W*0.12, W*0.52, W*0.22, W*0.14],
                      repeatRows=1)
            t.setStyle(_pdf_tbl_style(colours))
            story.append(t)

    # ── Disclaimer ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 14))
    story.append(_pdf_hr(colours))
    story.append(Paragraph(
        "DISCLAIMER: This report contains concept-level estimates (±40%) for comparative "
        "purposes only. CAPEX and OPEX figures are indicative AUD 2024 values and must not "
        "be used for procurement, funding approval, or investment decisions without "
        "site-specific engineering assessment by a qualified engineer. Energy and carbon "
        "values are based on standard industry parameters unless site calibration has been "
        "applied. All calculations follow Metcalf &amp; Eddy 5th Edition methodology.",
        styles["disc"]))

    doc.multiBuild(story)
    return buf.getvalue()
# ─────────────────────────────────────────────────────────────────────────────

def _docx_shade(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd elements first
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # OOXML schema: shd must appear before vAlign, hideMark, headers in tcPr
    # Correct order: tcW, gridSpan, vMerge, tcBorders, shd, noWrap, tcMar,
    #                textDirection, tcFitText, vAlign, hideMark
    insert_before = None
    for tag in (qn("w:noWrap"), qn("w:tcMar"), qn("w:textDirection"),
                qn("w:tcFitText"), qn("w:vAlign"), qn("w:hideMark")):
        el = tcPr.find(tag)
        if el is not None:
            insert_before = el
            break
    if insert_before is not None:
        tcPr.insert(list(tcPr).index(insert_before), shd)
    else:
        tcPr.append(shd)


def _docx_hr(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = p.paragraph_format.space_after = 0
    pPr = p._p.get_or_add_pPr()
    # Remove existing pBdr if any
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom)
    # Insert pBdr before ind/jc elements (correct schema position)
    insert_before = None
    for tag in (qn("w:shd"), qn("w:tabs"), qn("w:spacing"), qn("w:ind"), qn("w:jc")):
        el = pPr.find(tag)
        if el is not None:
            insert_before = el
            break
    if insert_before is not None:
        pPr.insert(list(pPr).index(insert_before), pBdr)
    else:
        pPr.append(pBdr)


def _docx_table(doc, data_rows, col_widths_cm, font_size=8):
    """data_rows[0] = header row (list of strings), rest = data."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    if not data_rows:
        return
    n_cols = len(data_rows[0])
    t = doc.add_table(rows=len(data_rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(data_rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "—")
            run.font.size = Pt(font_size)
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _docx_shade(cell, BLUE_HEX)
            elif r_idx % 2 == 0:
                _docx_shade(cell, GREY_LT)
    # Column widths
    for r in t.rows:
        for i, cell in enumerate(r.cells):
            if i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])


def _docx_heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x6A, 0xA5)
        h.runs[0].font.size = Pt(13 if level == 1 else 11)


def _docx_body(doc, text, size=9):
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(size)
    return p


def _docx_bullet(doc, text):
    from docx.shared import Pt
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9)


def _render_dict_table_docx(doc, data: dict, col_cm: list = None):
    hdrs = data.get("headers", [])
    rows = data.get("rows", [])
    if not hdrs or not rows:
        return
    n = len(hdrs)
    if col_cm is None:
        col_cm = [4.5] + [12.0 / (n-1)] * (n-1) if n > 1 else [16.0]
    tbl_data = [hdrs] + [[str(r.get(h, "—")) for h in hdrs] for r in rows]
    _docx_table(doc, tbl_data, col_cm)


def _render_list_table_docx(doc, data: list, col_cm: list = None):
    if not data:
        return
    hdrs = list(data[0].keys())
    n = len(hdrs)
    if col_cm is None:
        col_cm = [4.5] + [11.5 / (n-1)] * (n-1) if n > 1 else [16.0]
    tbl_data = [hdrs] + [[str(r.get(h, "—")) for h in hdrs] for r in data]
    _docx_table(doc, tbl_data, col_cm)


def _docx_base(report):
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)
    # Fix zoom element: add required w:percent attribute (OOXML schema compliance)
    settings_el = doc.settings.element
    zoom_els = settings_el.findall(qn("w:zoom"))
    for z in zoom_els:
        if not z.get(qn("w:percent")):
            z.set(qn("w:percent"), "100")
    return doc


# ─────────────────────────────────────────────────────────────────────────────
# EXECUTIVE SUMMARY WORD
# ─────────────────────────────────────────────────────────────────────────────

def _docx_executive(report: ReportObject) -> bytes:
    from docx.shared import Pt, RGBColor, Cm
    doc = _docx_base(report)

    # Title
    h = doc.add_heading(report.project_name or "Treatment Technology Comparison", 0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x6A, 0xA5)
    _docx_body(doc, f"{report.plant_name or ''} — Executive Summary", 11)
    _docx_body(doc,
        f"Prepared by: {report.prepared_by or 'ph2o consulting'}  |  "
        f"Date: {datetime.now().strftime('%B %Y')}  |  "
        f"Scenarios: {len(report.scenario_names)}", 8)
    _docx_hr(doc)

    # Overview
    _docx_heading(doc, "Study Overview")
    _docx_body(doc,
        f"This executive summary presents the results of a concept-stage treatment technology "
        f"comparison study for {report.plant_name or 'the facility'}. "
        f"{len(report.scenario_names)} treatment scenarios were evaluated across capital cost, "
        f"operating cost, lifecycle cost, energy, carbon, and risk. "
        f"All estimates are concept-level (±40%) for comparative purposes only.")

    # Comparison table
    if report.comparison_table:
        _docx_heading(doc, "Multi-Criteria Comparison")
        _docx_hr(doc)
        _render_list_table_docx(doc, report.comparison_table)
        _docx_body(doc, "Table 1: Multi-criteria comparison summary", 7)

    # Cost
    if report.cost_table:
        _docx_heading(doc, "Cost Summary")
        _docx_hr(doc)
        _render_dict_table_docx(doc, report.cost_table)
        _docx_body(doc,
            "All CAPEX in AUD 2024 (±40%). Lifecycle cost = OPEX + annualised CAPEX at 7% / 30 years.", 7)

    # Carbon
    if report.carbon_table:
        _docx_heading(doc, "Carbon & Energy")
        _docx_hr(doc)
        _render_dict_table_docx(doc, report.carbon_table)

    # Conclusions
    _docx_heading(doc, "Conclusions")
    _docx_hr(doc)
    if report.preferred_scenario:
        _docx_body(doc,
            f"The {report.preferred_scenario} scenario is identified as the preferred option "
            f"based on the multi-criteria assessment.")
    else:
        _docx_body(doc,
            "No preferred scenario nominated. Selection should consider site constraints, "
            "effluent standards, and stakeholder priorities.")

    # Disclaimer
    doc.add_paragraph()
    _docx_hr(doc)
    p = _docx_body(doc,
        "DISCLAIMER: Concept-level estimates (±40%). Not for procurement or funding approval. "
        "Prepared by ph2o consulting using the Water Utility Planning Platform.", 7)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# COMPREHENSIVE REPORT WORD
# ─────────────────────────────────────────────────────────────────────────────

def _docx_comprehensive(report: ReportObject) -> bytes:
    from docx.shared import Pt, RGBColor, Cm
    doc = _docx_base(report)

    # Cover
    h = doc.add_heading(report.project_name or "Treatment Technology Comparison", 0)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x6A, 0xA5)
    _docx_body(doc, report.plant_name or "Wastewater Treatment Facility", 12)
    doc.add_paragraph()
    meta = [
        ["Report Type:", "Concept-Stage Treatment Technology Comparison"],
        ["Scenarios:", str(len(report.scenario_names))],
        ["Prepared by:", report.prepared_by or "ph2o consulting"],
        ["Reviewed by:", report.reviewed_by or "—"],
        ["Date:", datetime.now().strftime("%d %B %Y")],
        ["Status:", "CONCEPT — ±40% ESTIMATE"],
    ]
    _docx_table(doc, meta, [4.5, 11.5])
    doc.add_paragraph()
    _docx_hr(doc)
    _docx_body(doc,
        "⚠ CONCEPT LEVEL ESTIMATE — ±40% — FOR COMPARATIVE PURPOSES ONLY. "
        "NOT FOR PROCUREMENT, FUNDING APPROVAL, OR INVESTMENT DECISIONS.", 8)
    doc.add_page_break()

    # 1. Introduction
    _docx_heading(doc, "1. Introduction")
    _docx_hr(doc)
    _docx_body(doc,
        f"This report presents the results of a concept-stage treatment technology comparison "
        f"study for {report.plant_name or 'the facility'}. The study applies first-principles "
        f"engineering calculations based on Metcalf & Eddy (5th Edition), WEF MOP 32/35, "
        f"and published technology-specific references.")
    _docx_body(doc,
        "All cost estimates are ±40% and are not suitable for procurement, detailed "
        "feasibility, or funding approval without site-specific investigation.")

    # 2. Scope
    _docx_heading(doc, "2. Study Scope and Methodology")
    _docx_hr(doc)
    _docx_body(doc, f"{len(report.scenario_names)} treatment scenarios were evaluated:")
    for name in (report.scenario_names or []):
        _docx_bullet(doc, name)
    _docx_body(doc, "Each scenario was assessed on:")
    criteria = [
        "Capital cost (CAPEX) — civil, mechanical, electrical, membrane works",
        "Operating cost (OPEX) — electricity, chemicals, sludge disposal",
        "Lifecycle cost — CAPEX annualised at 7% over 30 years plus OPEX",
        "Specific energy consumption (kWh/ML treated)",
        "Carbon footprint — Scope 1 (process) and Scope 2 (grid electricity)",
        "Effluent quality — BOD, TSS, TN, NH4, TP",
        "Biosolids production and sludge yield",
        "Technology risk — maturity, implementation, operational complexity, regulatory",
    ]
    for c in criteria:
        _docx_bullet(doc, c)

    # 3. Comparison
    if report.comparison_table:
        _docx_heading(doc, "3. Multi-Criteria Comparison")
        _docx_hr(doc)
        _docx_body(doc, "Table 1 presents key comparative metrics across all scenarios.")
        _render_list_table_docx(doc, report.comparison_table)
        _docx_body(doc, "Table 1: Multi-criteria comparison summary", 7)

    # 4b. Process Design Summary + PFDs
    _docx_heading(doc, "4. Process Design Summary")
    _docx_hr(doc)
    _docx_body(doc,
        "The following tables summarise key process design parameters for each scenario, "
        "derived from first-principles calculations.")

    design_data = getattr(report, "scenario_design_data", {})
    for scen_name in (report.scenario_names or []):
        sd = design_data.get(scen_name, {})
        tech_seq = sd.get("tech_sequence", [])
        tp_all   = sd.get("tech_performance", {})
        tech_code = tech_seq[0] if tech_seq else None
        perf = tp_all.get(tech_code, {}) if tech_code else {}

        if perf:
            _docx_heading(doc, scen_name, 2)
            def _fmt(v, unit=""):
                if not v: return "—"
                return f"{float(v):,.0f}{unit}"

            rows = [["Parameter", "Value"]]
            if perf.get("reactor_volume_m3"): rows.append(["Bioreactor volume", _fmt(perf["reactor_volume_m3"], " m³")])
            if perf.get("footprint_m2"): rows.append(["Process footprint", _fmt(perf["footprint_m2"], " m2")])
            if perf.get("hydraulic_retention_time_hr"): rows.append(["HRT", f"{perf['hydraulic_retention_time_hr']:.1f} hr"])
            mlss = perf.get("mlss_mg_l") or perf.get("mlss_granular_mg_l")
            if mlss: rows.append(["MLSS", _fmt(mlss, " mg/L")])
            if perf.get("o2_demand_kg_day"): rows.append(["O₂ demand", _fmt(perf["o2_demand_kg_day"], " kg/day")])
            if perf.get("sludge_production_kgds_day"): rows.append(["Sludge production", _fmt(perf["sludge_production_kgds_day"], " kgDS/day")])
            eff = []
            for k, l in [("effluent_bod_mg_l","BOD"),("effluent_tss_mg_l","TSS"),("effluent_tn_mg_l","TN"),("effluent_nh4_mg_l","NH4")]:
                if perf.get(k) is not None: eff.append(f"{l} {perf[k]:.0f}")
            if eff: rows.append(["Effluent quality", "  ".join(eff) + " mg/L"])
            kwh = sd.get("specific_energy_kwh_kl", 0)
            if kwh: rows.append(["Specific energy", f"{kwh*1000:.0f} kWh/ML"])
            _docx_table(doc, rows, [7.0, 9.0])
            _docx_body(doc, f"Note: Process flow diagram for {scen_name} — see below.", 7)

    _docx_body(doc,
        "Process flow diagrams are included in the PDF version of this report. "
        "The Word version contains the design parameter tables above for each scenario.",
        7)

    # 5. Cost
    if report.cost_table:
        _docx_heading(doc, "5. Capital and Lifecycle Cost Assessment")
        _docx_hr(doc)
        _docx_body(doc,
            "Table 3 summarises capital cost, operating cost, and lifecycle cost. "
            "All costs are AUD 2024, concept-level (±40%).")
        _render_dict_table_docx(doc, report.cost_table)
        _docx_body(doc, "Table 3: Cost summary (AUD 2024, ±40%)", 7)
        _docx_body(doc,
            "CAPEX includes bioreactor civil, mechanical, electrical, membranes (where applicable), "
            "secondary clarifiers, and instrumentation. Excludes land, site preparation, headworks, "
            "sludge treatment, buildings, and owner's costs.")

    # 5a. OPEX Breakdown
    if getattr(report, "opex_breakdown_table", None):
        _docx_heading(doc, "5a. OPEX Cost Drivers")
        _docx_hr(doc)
        _docx_body(doc,
            "Table 3a breaks down annual operating cost by category to identify the primary "
            "drivers of the OPEX differential between scenarios. Energy and sludge disposal "
            "are typically the two largest components for biological treatment processes.")
        _render_dict_table_docx(doc, report.opex_breakdown_table)
        _docx_body(doc,
            "Table 3a: OPEX breakdown by category (AUD 2024/yr). "
            "Percentages shown as proportion of total annual OPEX.", 7)

    # 5b. Specific Performance Metrics
    if getattr(report, "specific_metrics_table", None):
        _docx_heading(doc, "5b. Specific Performance Metrics")
        _docx_hr(doc)
        _docx_body(doc,
            "Table 3b presents normalised performance metrics for benchmarking and "
            "comparison. Specific footprint (m2/MLD) and specific sludge (kgDS/ML) "
            "are standard Australian utility capital planning metrics. "
            "Carbon intensity (kgCO2e/kL) enables comparison with water supply benchmarks.")
        _render_dict_table_docx(doc, report.specific_metrics_table)
        _docx_body(doc, "Table 3b: Specific performance metrics.", 7)

    # 6. Energy & Carbon
    if report.carbon_table:
        _docx_heading(doc, "6. Energy and Carbon Footprint")
        _docx_hr(doc)
        _docx_body(doc,
            "Table 4 summarises carbon emissions. Scope 1 includes N2O from biological "
            "nitrogen removal and CH4 from sludge handling. Scope 2 reflects grid electricity "
            "consumption. Avoided emissions reflect CHP electricity generation where applicable.")
        _render_dict_table_docx(doc, report.carbon_table)
        _docx_body(doc, "Table 4: Carbon and energy summary", 7)

    # 7. Risk
    if report.risk_table:
        _docx_heading(doc, "7. Risk Assessment")
        _docx_hr(doc)
        _docx_body(doc,
            "Table 5 presents technology risk across technical, implementation, operational, "
            "and regulatory dimensions. Scores are indicative based on technology maturity "
            "and reference plant availability.")
        _render_dict_table_docx(doc, report.risk_table)
        _docx_body(doc, "Table 5: Technology risk assessment", 7)

    # 8. Decision Framework (if decision engine ran)
    dec = getattr(report, "decision", None)
    if dec and getattr(dec, "recommended_tech", None):
        doc.add_page_break()
        _docx_heading(doc, "8. Decision Framework")
        _docx_hr(doc)

        # Selection basis banner
        _docx_body(doc, f"Selection basis: {dec.selection_basis}", 9)
        doc.add_paragraph()

        # Recommendation
        _docx_heading(doc, "Recommended Option", 2)
        _docx_body(doc, getattr(dec, "display_recommended_label", dec.recommended_label), 11)
        for reason in dec.why_recommended:
            _docx_bullet(doc, reason)

        # Non-viable
        if dec.non_viable:
            _docx_heading(doc, "Base Case Non-Compliant Options (without intervention)", 2)
            for nv in dec.non_viable:
                _docx_bullet(doc, f"{nv} — non-compliant as base case; compliant with engineering intervention")

        # Regulatory note
        reg_note = getattr(dec, "regulatory_note", "")
        if reg_note:
            _docx_heading(doc, "Regulatory Note", 2)
            _docx_body(doc, reg_note)

        # Alternative pathways
        alt_paths = getattr(dec, "alternative_pathways", [])
        if alt_paths:
            _docx_heading(doc, "Alternative Pathways — Interventions for Non-Compliant Base Cases", 2)
            _docx_body(doc,
                "The following engineering interventions can make currently non-compliant "
                "options viable. Evaluate before committing to procurement.")
            for p in alt_paths:
                icon = "✓" if p.achieves_compliance else "⚠"
                _docx_heading(doc, f"{p.tech_label} + Intervention", 3)
                rows = [
                    ["Intervention",        p.intervention],
                    ["CAPEX increment",     f"+${p.capex_delta_m:.1f}M"],
                    ["OPEX increment",      f"+${p.opex_delta_k:.0f}k/yr"],
                    ["Total LCC (pathway)", f"${p.lcc_total_k:.0f}k/yr"],
                    ["Achieves compliance", f"{icon} {'Yes' if p.achieves_compliance else 'Partial'}"],
                    ["Procurement",         p.procurement],
                    ["Regulatory",          p.regulatory],
                ]
                _docx_table(doc, rows, [5.0, 11.0])
                _docx_body(doc, p.summary, 8)
                if p.residual_risks:
                    _docx_body(doc, "Residual risks:", 9)
                    for r in p.residual_risks:
                        _docx_bullet(doc, r)

        # Client decision framing
        cf = getattr(dec, "client_framing", None)
        if cf:
            _docx_heading(doc, "Client Decision Framing", 2)
            _docx_body(doc,
                "The client is not choosing between technologies. The client is choosing "
                "how to manage risk, capital expenditure, and delivery complexity.")
            rows = [
                ["", cf.option_a_label, cf.option_b_label],
                ["What you get",    " | ".join(cf.option_a_bullets[:3]),
                                    " | ".join(cf.option_b_bullets[:3])],
                ["Risks accepted",  " | ".join(cf.option_a_risks[:2]),
                                    " | ".join(cf.option_b_risks[:2])],
            ]
            _docx_table(doc, rows, [3.5, 6.5, 6.0])
            _docx_body(doc, "Decision depends on:")
            for f in cf.deciding_factors:
                _docx_bullet(doc, f)
            _docx_body(doc, cf.framing_note, 8)

        # Technology profiles
        if dec.profiles:
            _docx_heading(doc, "Technology Delivery & Constructability", 2)
            _docx_body(doc,
                "Procurement model and constructability assessment for each technology.")
            hdr = ["Technology", "D&C", "DBOM", "Alliance", "Constructability", "Recommended"]
            rows = [hdr]
            for name, profile in dec.profiles.items():
                rows.append([
                    name,
                    profile.delivery.dnc.value,
                    profile.delivery.dbom.value,
                    profile.delivery.alliance.value,
                    profile.constructability.overall.value,
                    profile.delivery.recommended_model,
                ])
            _docx_table(doc, rows, [4.0, 2.0, 2.0, 2.0, 3.0, 3.0])

            # Failure modes
            _docx_heading(doc, "Key Failure Modes", 2)
            for name, profile in dec.profiles.items():
                _docx_heading(doc, name, 3)
                _docx_body(doc, f"Critical: {profile.failure_modes.critical_note}", 9)
                for fm in profile.failure_modes.modes:
                    _docx_bullet(doc,
                        f"{fm.name} — Likelihood: {fm.likelihood}, "
                        f"Consequence: {fm.consequence}. {fm.mitigation}")

        # Strategic Insight (process intensification vs robustness framing)
        si = getattr(dec, "strategic_insight", "")
        if si:
            _docx_heading(doc, "Strategic Insight", 2)
            _docx_body(doc,
                "This framing should guide technology selection and long-term asset strategy.")
            # Split on double-newline paragraphs
            for para in si.split("\n\n"):
                para = para.strip()
                if para:
                    _docx_body(doc, para)

        # Recommended Approach (parallel evaluation steps)
        ra = getattr(dec, "recommended_approach", [])
        if ra:
            _docx_heading(doc, "Recommended Approach", 2)
            _docx_body(doc,
                "No premature technology lock-in. Parallel concept design "
                "validation before final selection.")
            for step in ra:
                if step.startswith("  "):
                    # Sub-item — render as indented body text
                    _docx_body(doc, f"→ {step.strip()}", 9)
                elif step.endswith(":"):
                    _docx_body(doc, step, 9)
                else:
                    _docx_bullet(doc, step)

        # Trade-offs
        if dec.trade_offs:
            _docx_heading(doc, "Trade-off Summary", 2)
            for t in dec.trade_offs:
                _docx_bullet(doc, t)

        # Financial Risk Perspective
        alt_paths_fr = getattr(dec, "alternative_pathways", [])
        rec_capex = recommended_capex_for_report  if 'recommended_capex_for_report' in dir() else None
        _docx_heading(doc, "Financial Risk Perspective", 2)
        _docx_body(doc,
            "Capital and operating cost structures carry different financial risk profiles "
            "over the asset lifecycle. Both pathway characteristics are summarised below.")
        # Build rows from decision data
        cf_fr = getattr(dec, "client_framing", None)
        if cf_fr and alt_paths_fr:
            a_fr = alt_paths_fr[0]
            opt_a_capex = next((b for b in cf_fr.option_a_bullets if "CAPEX" in b), "—")
            opt_b_capex = next((b for b in cf_fr.option_b_bullets if "CAPEX" in b), "—")
            rows_fr = [
                ["Risk dimension",
                 f"Option A: {dec.recommended_label}",
                 f"Option B: {a_fr.tech_label} + thermal"],
                ["CAPEX exposure",
                 opt_a_capex.replace("CAPEX: ", ""),
                 opt_b_capex.replace("CAPEX: ", "")],
                ["OPEX exposure",
                 "Moderate — electricity (blowers) + vendor DBOM fee",
                 "Higher — heating energy + methanol (ongoing chemical cost)"],
                ["Energy dependency",
                 "Moderate — MABR gas-side pressure control",
                 "High — continuous heating (critical: 15°C compliance threshold)"],
                ["Chemical dependency",
                 "None — biological process only",
                 "High — methanol supply security, storage, dosing"],
                ["LCC sensitivity",
                 "Sensitive to electricity price, MABR vendor fee",
                 "Sensitive to gas/heating energy price, methanol price"],
                ["Long-term financial risk",
                 "Technology obsolescence if MABR market contracts; "
                 "higher ongoing vendor dependency",
                 "Commodity price exposure (methanol, energy); "
                 "heating system capital replacement at 15–20yr"],
            ]
            _docx_table(doc, rows_fr, [4.0, 6.0, 6.0])
            _docx_body(doc,
                "Both options carry lifecycle cost uncertainty of ±25%. "
                "BNR+thermal OPEX is more exposed to energy and chemical price inflation. "
                "MABR LCC is more exposed to vendor contract pricing and technology evolution.",
                8)

        # Confidence
        conf = getattr(dec, "confidence", None)
        if conf:
            _docx_heading(doc, f"Recommendation Confidence: {conf.level}", 2)
            for d in conf.drivers:
                _docx_bullet(doc, d)
            if conf.caveats:
                _docx_body(doc, "Caveats:")
                for c in conf.caveats:
                    _docx_bullet(doc, c)

    # 9. Conclusions
    _docx_heading(doc, "9. Conclusions and Recommendations")
    _docx_hr(doc)
    if dec and getattr(dec, "recommended_tech", None):
        _docx_body(doc, dec.conclusion)
    elif report.preferred_scenario:
        _docx_body(doc,
            f"The {report.preferred_scenario} scenario is identified as the preferred option, "
            f"offering the best balance of lifecycle cost, energy performance, carbon footprint, and risk.")
    else:
        _docx_body(doc,
            "Selection of a preferred option should be informed by site-specific constraints, "
            "effluent standards, utility priorities, and procurement strategy.")
    recs = [
        "Site constraints — footprint, existing infrastructure, upgrade pathway",
        "Effluent quality requirements and regulatory licence conditions",
        "Utility priorities — energy, carbon, cost reduction",
        "Local contractor capability and procurement strategy",
        "Detailed geotechnical and hydraulic assessment",
    ]
    for r in recs:
        _docx_bullet(doc, r)
    _docx_body(doc,
        "A detailed feasibility study with site-specific cost estimation (±15-20%) "
        "is recommended before proceeding to detailed design or procurement.")

    # Closing action statement — programme urgency
    dec_close = getattr(report, "decision", None)
    if dec_close and getattr(dec_close, "recommended_approach", []):
        from docx.shared import Pt
        p = doc.add_paragraph()
        run = p.add_run(
            "Recommended next step: initiate parallel concept validation immediately "
            "to maintain programme schedule and avoid premature technology lock-in."
        )
        run.bold = True
        run.font.size = Pt(9)

    # Appendix — Assumptions
    if report.assumptions_appendix:
        doc.add_page_break()
        _docx_heading(doc, "Appendix A — Key Assumptions")
        _docx_hr(doc)
        skip_prefixes = ["capex_unit_costs", "pump_per_kw", "blower_per_kw",
                         "vsd_per_kw", "instrumentation", "design_contingency",
                         "contractor_margin", "client_oncosts", "return_sludge",
                         "gravity_thickener", "belt_thickener", "centrifuge",
                         "struvite", "fine_screen", "mbr_membrane_flat"]
        clean = []
        for item in report.assumptions_appendix:
            param = str(item.get("Parameter", ""))
            if any(param.startswith(p) for p in skip_prefixes):
                continue
            clean.append([
                item.get("Category", ""),
                _format_assumption_param(param),
                str(item.get("Value", "")),
                item.get("User Override", ""),
            ])
        if clean:
            _docx_table(doc, [["Category", "Parameter", "Value", "Override"]] + clean,
                        [2.5, 8.0, 3.5, 2.0])

    # Disclaimer
    doc.add_paragraph()
    _docx_hr(doc)
    _docx_body(doc,
        "DISCLAIMER: This report contains concept-level estimates (±40%) for comparative "
        "purposes only. CAPEX and OPEX figures are indicative AUD 2024 values and must not "
        "be used for procurement, funding approval, or investment decisions without "
        "site-specific engineering assessment by a qualified engineer.", 7)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
